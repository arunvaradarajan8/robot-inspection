from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path('docs/Infrastructure_Defect_Detection_Application_Overview.docx')
FOOTER_TEXT = 'Arun Varadarajan | arunvaradarajan@txdot.gov'

ORANGE = RGBColor(221, 96, 0)
LIGHT_ORANGE = RGBColor(238, 148, 72)
DARK_BLUE = RGBColor(20, 65, 96)
BODY = RGBColor(35, 35, 35)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    tc_pr.append(shading)


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = FOOTER_TEXT
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(90, 90, 90)


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    add_footer(section)

    styles = document.styles
    normal = styles['Normal']
    normal.font.name = 'Aptos'
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BODY

    title = styles['Title']
    title.font.name = 'Aptos Display'
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = ORANGE

    h1 = styles['Heading 1']
    h1.font.name = 'Aptos Display'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = ORANGE

    h2 = styles['Heading 2']
    h2.font.name = 'Aptos Display'
    h2.font.size = Pt(13.5)
    h2.font.bold = True
    h2.font.color.rgb = LIGHT_ORANGE

    h3 = styles['Heading 3']
    h3.font.name = 'Aptos'
    h3.font.size = Pt(11.5)
    h3.font.bold = True
    h3.font.color.rgb = DARK_BLUE


def para(document, text, style=None):
    p = document.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    p.add_run(text)
    return p


def bullets(document, items):
    for item in items:
        p = document.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def numbered(document, items):
    for item in items:
        p = document.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def code_block(document, text):
    table = document.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'F4F6F8')
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.8)
    run.font.color.rgb = RGBColor(45, 45, 45)
    document.add_paragraph()


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_shading(hdr[i], 'F4B183')
        run = hdr[i].paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    document.add_paragraph()


def new_chapter(document, title):
    if len(document.paragraphs) > 2:
        document.add_page_break()
    document.add_heading(title, level=1)


def add_title_page(document):
    title = document.add_paragraph(style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run('Infrastructure Defect Detection\nDigital Twin Inspection System')

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        'Python Implementation, ROS 2 Node Design, Mission Control APIs, '
        'OAK-D Fusion, Trimble X7 Scan Coordination, Digital Twin State, '
        'and Spot Navigation Control'
    )
    run.font.size = Pt(15)
    run.font.color.rgb = DARK_BLUE

    para(
        document,
        'Prepared for field deployment planning and system handoff. This document '
        'summarizes the end-to-end inspection application, the hardware stack, '
        'the ROS 2 nodes, the control-panel workflow, the data products, and the '
        'operational considerations for autonomous infrastructure inspection.',
    )
    add_table(
        document,
        ['Document Attribute', 'Value'],
        [
            ['Author', 'Arun Varadarajan'],
            ['Contact', 'arunvaradarajan@txdot.gov'],
            ['Primary Compute', 'NVIDIA Jetson Orin Nano'],
            ['Robot Platform', 'Boston Dynamics Spot'],
            ['Perception Camera', 'Luxonis OAK-D-W / OAK-D Pro W class RGB-D camera'],
            ['High-Fidelity Scanner', 'Trimble X7'],
            ['Control Host', 'Windows tablet or Windows laptop running Trimble Perspective bridge role'],
        ],
    )


CHAPTERS = [
    (
        '1. Executive System Overview',
        [
            (
                'Inspection Mission Concept',
                'The application is designed as a robotic infrastructure inspection '
                'pipeline that combines live AI perception, robot movement, high-fidelity '
                'laser scanning, and digital-twin marker persistence. Spot carries the '
                'Jetson compute payload and OAK-D camera, moves between inspection stations, '
                'and pauses when a detailed Trimble X7 scan is needed. The software chooses '
                'scan points based on known defect markers, map frontiers, and scan gating '
                'rules so the robot avoids unnecessary scans when detections are absent or '
                'low-confidence.',
            ),
            (
                'Primary Data Flow',
                'The field pipeline begins with live RGB imagery from the OAK camera. YOLO '
                'turns images into 2D defect detections, aligned depth converts detections '
                'into 3D positions, and the digital-twin nodes persist those detections as '
                'markers and rescan goals. Trimble X7 scans are ingested as LAS or LAZ files, '
                'reduced for transfer when needed, and converted into ROS point clouds and '
                'occupancy maps. The planner uses those maps and markers to issue inspection '
                'goals to Nav2, an HTTP bridge, or the Spot SDK backend.',
            ),
            (
                'Operational Outcome',
                'The intended output is a compact digital-twin package containing defect '
                'markers, coordinate anchoring information, mission status, and processed '
                'scan products. The raw Trimble scans can remain on the Perspective control '
                'host, while the Jetson keeps smaller mission artifacts and map products for '
                'navigation and field decision-making.',
            ),
        ],
    ),
    (
        '2. Hardware Inventory',
        [
            (
                'Boston Dynamics Spot',
                'Spot is the mobile robot platform that physically carries the sensing and '
                'compute payload. The application treats Spot as the motion platform while '
                'retaining Spot internal stability and locomotion control. Higher-level '
                'inspection goals are produced by ROS and then handed to a movement backend.',
            ),
            (
                'Spot GXP Payload Interface',
                'The Spot GXP provides payload mounting and power/communications breakout. '
                'For this system it is treated as the integration point for the Jetson, local '
                'networking hardware, camera wiring, and sealed cable management. It can '
                'provide regulated outputs including 5 V, 12 V, and 24 V rails depending on '
                'the payload wiring choices.',
            ),
            (
                'Jetson Orin Nano',
                'The Jetson is the main field compute target. It runs ROS 2 Jazzy, the OAK '
                'camera driver stack, YOLO inference, digital-twin processing, scan watchers, '
                'planners, and robot goal bridge nodes. It should be powered through its '
                'barrel jack using an appropriate regulated DC-DC converter.',
            ),
            (
                'Luxonis OAK-D-W / OAK-D Pro W Class Camera',
                'The OAK camera supplies RGB imagery, stereo depth, camera calibration, and '
                'optionally odometry or VSLAM data through DepthAI ROS. The center color '
                'camera is used for AI detection, while the stereo pair and IR capability '
                'support depth and localization. The camera is mounted from the back using '
                'its M4/VESA-style mounting points; the front optical face needs a generous '
                'unobstructed opening.',
            ),
            (
                'Trimble X7',
                'The Trimble X7 provides high-fidelity laser scans. The current software '
                'does not directly command the scanner hardware at the ROS node level; it '
                'coordinates with the Perspective control host, then ingests completed LAS, '
                'LAZ, or E57 exports. The scanner is expected to operate while the robot is '
                'still at a station.',
            ),
            (
                'Control Host',
                'The control host is a Windows tablet or Windows laptop running Trimble '
                'Perspective and the Tkinter bridge app. It receives scan requests from the '
                'Jetson, watches the Perspective export folder, transfers scans, and downloads '
                'compact digital-twin artifacts at mission end.',
            ),
            (
                'Network and Power Hardware',
                'A field router or private mission LAN connects the Jetson, tablet/laptop, '
                'and any Spot command services. A USB-C powered router is useful when Ethernet '
                'is not practical. The system also needs USB cables for the OAK, barrel-jack '
                'power for the Jetson, cable retention, strain relief, and appropriate fusing '
                'or DC-DC regulation.',
            ),
        ],
    ),
    (
        '3. Mechanical Integration',
        [
            (
                'OAK Camera Case',
                'The OAK case should secure the camera from the rear mounting points rather '
                'than clamping the front lens face. A practical 3D-printed enclosure uses '
                'approximately 3 mm walls, 4-6 mm thickness around the rear mounting area, '
                'and a large chamfered optical window across the front. The front should not '
                'have tight individual holes around the lenses because the wide-FOV stereo '
                'cameras can vignette or reflect IR from nearby surfaces.',
            ),
            (
                'Ventilation',
                'The OAK electronics produce heat during operation. Top exhaust slots and '
                'side intake slots reduce thermal buildup. Top vent slots should be shielded '
                'or louvered for field use, while the bottom should remain mostly closed to '
                'reduce dust and splash exposure. Vents should avoid lens and IR emitter paths.',
            ),
            (
                'Trimble X7 Mounting',
                'The X7 must remain stable during scans. The robot should stop before the '
                'scanner begins a high-fidelity station scan. Any X7 mount should minimize '
                'vibration, preserve scanner line of sight, and avoid collisions with the '
                'OAK field of view, Spot body, or lighting hardware.',
            ),
            (
                'Lighting',
                'The Ulanzi 40 W Pro LED can improve RGB detection in low-light areas. It '
                'should be mounted so it lights the inspected surface without shining directly '
                'into the OAK stereo pair or Trimble scanner. IR from the OAK and visible LED '
                'light should be tested together for exposure and depth stability.',
            ),
        ],
    ),
    (
        '4. Power Architecture',
        [
            (
                'Jetson Power',
                'The Jetson should be powered through the barrel jack. Use a regulated DC-DC '
                'converter sized for the Jetson load and startup current. The converter should '
                'be mounted securely, fused, and strain-relieved.',
            ),
            (
                'Router Power',
                'A compact router powered by USB-C can share the Spot/Jetson mission network '
                'with the tablet or laptop. Power it from a regulated 5 V USB-C supply or '
                'appropriate DC-DC converter. Avoid brownouts during robot motion or scanner '
                'startup.',
            ),
            (
                'Camera Power',
                'The OAK camera should use its own regulated 5 V USB-C power path when the '
                'Jetson USB port is not intended to carry the camera power budget. Keep OAK '
                'power separate from USB data, use a short retained data cable to the Jetson, '
                'and verify the camera remains connected during robot motion.',
            ),
            (
                'Trimble Power',
                'The Trimble X7 is not expected to be powered from Spot. It should use its '
                'own battery or supported power arrangement. The integration assumes the '
                'scanner and Perspective host manage scan power independently.',
            ),
        ],
    ),
    (
        '5. Network Architecture',
        [
            (
                'Mission LAN',
                'The Jetson, control host, and any Spot command endpoint must be reachable '
                'over the same private network. The system does not require internet access '
                'during field operation once dependencies and models are installed. A field '
                'router provides stable addressing and avoids dependency on external Wi-Fi.',
            ),
            (
                'ROS and HTTP Boundaries',
                'ROS 2 nodes run primarily on the Jetson. The control host communicates with '
                'the Jetson through SSH, SCP, HTTP endpoints, or shared folders. The HTTP '
                'bridge endpoints are intentionally simple so the Windows app can remain a '
                'thin coordination layer rather than a ROS machine.',
            ),
            (
                'Suggested Addressing',
                'Use static DHCP reservations or documented static IPs for the Jetson, control '
                'host, and Spot. The TRIMBLE_WINDOWS_URL field currently represents the '
                'Windows Perspective bridge URL that the Jetson calls for status and scan '
                'coordination.',
            ),
        ],
    ),
    (
        '6. Software Packages',
        [
            (
                'defect_detection Package',
                'This ROS package contains the AI detector, RGB-D fusion, scan decision logic, '
                'digital-twin nodes, planner nodes, visualization nodes, and robot command '
                'bridge. It is the main application package and is launched through the field '
                'scripts or full pipeline launch file.',
            ),
            (
                'pointcloud_bridge Package',
                'This package normalizes incoming ROS point clouds. It validates x, y, and z '
                'fields, enforces timestamp policy, optionally changes the output frame ID, '
                'and republishes the result to the configured point-cloud topic. It is generic '
                'and no longer tied to any legacy Spot EAP naming.',
            ),
            (
                'Perspective Bridge Tools',
                'The tools/trimble_perspective_bridge directory contains the current Windows '
                'control app. It establishes the mission control UI, SSHs into the Jetson, '
                'starts ROS, listens for scan requests, watches the scan export directory, '
                'and downloads final artifacts. The app is intended to run on the Windows '
                'tablet or laptop that also runs Trimble Perspective.',
            ),
        ],
    ),
    (
        '7. OAK-D Perception Pipeline',
        [
            (
                'RGB Input',
                'The image topic is configured as /oak/rgb/image_raw in the field environment. '
                'YOLO subscribes to that stream and publishes Detection2DArray messages on '
                '/detections_2d. The center RGB sensor is the primary source for visible '
                'defect classification.',
            ),
            (
                'Depth Input',
                'The OAK depth fusion node subscribes to /oak/rgb/depth and /oak/rgb/camera_info. '
                'The depth image must be aligned to the image used by YOLO. If the DepthAI ROS '
                'topic names differ, config/field.env should be updated after checking ros2 '
                'topic list on the Jetson.',
            ),
            (
                '3D Projection',
                'The fusion node takes each high-confidence 2D detection, samples robust depth '
                'inside the detection box, and projects the detection center into camera '
                'coordinates using the camera intrinsic matrix. The result is published as '
                'Detection3DArray on /detections_3d.',
            ),
        ],
    ),
    (
        '8. OAK Localization',
        [
            (
                'Odometry Source',
                'The system expects an OAK odometry or VSLAM source on /oak/odom. The '
                'oak_localization_bridge republishes that odometry into TF, commonly as '
                'oak_odom -> body. This is the mission localization source for planning and '
                'arrival checks.',
            ),
            (
                'Frame Strategy',
                'The robot body frame is configured as body in the field environment. The '
                'digital twin map frame is anchored to the OAK world frame after the first '
                'Trimble reference scan. This lets future defect markers and scan stations be '
                'computed relative to the camera-localized robot path.',
            ),
            (
                'Limitations',
                'Visual odometry can drift in low texture, heavy dust, fast motion, or poor '
                'lighting. The system uses it because it provides a useful local mission frame, '
                'but the Trimble scan and manual validation remain important for high-fidelity '
                'deliverables.',
            ),
        ],
    ),
    (
        '9. AI Detection',
        [
            (
                'YOLO Detector',
                'The yolo_detector node loads a YOLO engine and dataset YAML, subscribes to '
                'the RGB image topic, and publishes 2D detections. The dataset file defines '
                'the defect classes. The engine path is configured in config/field.env.',
            ),
            (
                'Confidence Gating',
                'Detections below the configured threshold do not trigger downstream scan '
                'requests. This keeps the X7 scan workflow efficient and prevents the robot '
                'from spending time scanning empty or uncertain areas.',
            ),
            (
                'Detection Artifacts',
                '2D detections are used for immediate scan decisions and camera preview. 3D '
                'detections become digital-twin markers, rescan goals, and visualization '
                'markers for RViz or mission-control display.',
            ),
        ],
    ),
    (
        '10. Trimble X7 Scan Workflow',
        [
            (
                'Control Model',
                'The ROS stack does not directly query or command the X7 hardware. Instead, '
                'the Perspective control host handles the scanner session and produces export '
                'files. The Jetson bridge requests scans and waits for exported LAS, LAZ, or '
                'E57 data to become available.',
            ),
            (
                'Initial Reference Scan',
                'A reference scan can be requested on startup even if no defects are visible. '
                'That scan defines the first high-fidelity map context and allows the frame '
                'anchor node to relate the Trimble map frame to the OAK odometry frame.',
            ),
            (
                'Scan Efficiency',
                'After the reference scan, additional X7 scans are gated by high-confidence '
                'AI detections, known defect rescan goals, frontier goals, and waypoint arrival '
                'events. This reduces the number of stationary scans and improves mission time.',
            ),
        ],
    ),
    (
        '11. Scan File Ingestion',
        [
            (
                'Trimble Scan Watcher',
                'trimble_scan_watcher.py watches a configured directory for stable completed '
                'scan files. It waits until files are old enough to be considered complete, '
                'then converts scan points into sensor_msgs/PointCloud2 on /trimble/x7/scan_points.',
            ),
            (
                'LAS and LAZ Handling',
                'The watcher uses laspy and lazrs for LAS/LAZ support. To keep wireless transfer '
                'reasonable, the control host can downsample or cap point count before sending '
                'the file to the Jetson.',
            ),
            (
                'Scan Request Requirement',
                'When require_scan_request is true, the watcher ignores random files until the '
                'scan decision logic requests a scan. This protects the mission from stale '
                'test files or accidental exports.',
            ),
        ],
    ),
    (
        '12. Digital Twin Map Builder',
        [
            (
                'Point Cloud to Occupancy',
                'pointcloud_to_occupancy.py converts the Trimble scan cloud into a 2D occupancy '
                'grid. It filters points by height and range, applies map padding, and publishes '
                'the resulting map on /digital_twin/map.',
            ),
            (
                'Map Usage',
                'The occupancy map is not a final survey deliverable. It is a navigation and '
                'planning abstraction that lets the robot choose inspection stations and frontier '
                'goals. The full-resolution Trimble scan remains the high-fidelity asset.',
            ),
            (
                'Resolution and Range',
                'The default map resolution is 0.10 m. Max range and z-filter settings should '
                'be tuned per site so floor noise, ceiling structure, and irrelevant distant '
                'points do not dominate the planning map.',
            ),
        ],
    ),
    (
        '13. Digital Twin Defect Markers',
        [
            (
                'Marker Persistence',
                'defect_map_node.py merges repeated detections, stores defect observations in '
                '/tmp/digital_twin_defects.yaml, and republishes markers for visualization. '
                'This persistence makes defects available after the original camera frame has '
                'passed.',
            ),
            (
                'Rescan Goals',
                'The same node publishes /digital_twin/rescan_goals. These are inspection '
                'targets that the infrastructure planner can prioritize before exploring new '
                'frontiers.',
            ),
            (
                'Marker Semantics',
                'Each marker carries class, confidence, pose, and observation count. The markers '
                'are AI-generated aids for inspection and should be validated against Trimble '
                'data, imagery, and human review before final reporting.',
            ),
        ],
    ),
    (
        '14. Frame Anchoring',
        [
            (
                'Anchor Purpose',
                'The frame anchor node stores the relationship between the digital twin map '
                'frame and the robot world frame. Without this step, the robot could create '
                'a scan map but would not know how to relate scan coordinates to future motion.',
            ),
            (
                'Anchor Timing',
                'The recommended field behavior is to anchor on the first completed Trimble '
                'reference scan. At that moment, the robot pose from OAK localization is captured '
                'and used as the transform between map and robot world.',
            ),
            (
                'Persistent Output',
                'The anchor is written to /tmp/digital_twin_anchor.yaml so the final twin package '
                'can include how the robot-localized mission related to the scan frame.',
            ),
        ],
    ),
    (
        '15. Planning Logic',
        [
            (
                'Infrastructure Planner',
                'infrastructure_planner.py chooses the next inspection goal. It prefers known '
                'defect rescan goals when available, then falls back to exploration or frontier '
                'goals from the map. It publishes PoseStamped goals on /infrastructure/inspection_goal.',
            ),
            (
                'Frontier Planner',
                'frontier_planner.py identifies map edges between known and unknown space. This '
                'lets the robot move to the boundary of explored space without relying solely '
                'on pre-scripted waypoints.',
            ),
            (
                'Cooldowns and Priorities',
                'Cooldown settings prevent rapid reissuing of goals. Priority configuration lets '
                'defect classes, confidence, distance, and map context influence which station '
                'should be visited next.',
            ),
        ],
    ),
    (
        '16. Robot Motion Backends',
        [
            (
                'Dry Run Backend',
                'The dry_run backend simulates waypoint arrival without commanding hardware. It '
                'is useful for UI testing, scan loop testing, and validating digital-twin logic '
                'before allowing robot motion.',
            ),
            (
                'Nav2 Backend',
                'The nav2 backend sends NavigateToPose goals to a Nav2 action server. This is '
                'useful if the robot platform provides a Nav2-compatible motion stack or if a '
                'simulation is being used.',
            ),
            (
                'HTTP Backend',
                'The HTTP backend posts goals to an external command service. This is useful '
                'when Spot SDK control is isolated into another process, machine, or safety layer.',
            ),
            (
                'Spot SDK Backend',
                'The spot_sdk backend connects to Spot, authenticates, acquires a lease, optionally '
                'powers on and stands the robot, sends an SE2 trajectory command, and waits for '
                'trajectory completion before publishing waypoint arrival.',
            ),
        ],
    ),
    (
        '17. Arrival Verification',
        [
            (
                'Why Arrival Matters',
                'The scan loop must not request a Trimble scan merely because a goal was sent. '
                'It should request a scan after the robot has actually reached and remained near '
                'the target pose. This protects scan quality and reduces wasted station scans.',
            ),
            (
                'TF-Based Check',
                'robot_goal_bridge.py can use TF to compare the active goal to the current base '
                'frame pose. It checks position tolerance, yaw tolerance, and stable duration. '
                'Only after the pose remains stable does it publish waypoint arrival.',
            ),
            (
                'Fallback Behavior',
                'If TF arrival checking is disabled or unavailable, backend completion can still '
                'act as a basic arrival signal. Field deployment should prefer TF arrival checks '
                'with the OAK localization source.',
            ),
        ],
    ),
    (
        '18. Perspective Control Host',
        [
            (
                'Windows App',
                'The current windows_app.py implementation provides a desktop Tkinter mission '
                'control app. It starts and stops ROS on the Jetson, exposes HTTP endpoints, '
                'shows mission state, watches scan exports, and downloads final artifacts.',
            ),
            (
                'Windows Tablet Deployment',
                'A Windows tablet is treated the same as a Windows laptop: install Python, run '
                'the dependency installer, launch the bridge app, and keep Trimble Perspective '
                'open on the same machine. The Jetson reaches the bridge through the mission LAN.',
            ),
            (
                'Automation Caveat',
                'If Perspective has no supported command-line or API trigger, the bridge still '
                'coordinates the workflow by showing scan requests, watching the export folder, '
                'and transferring completed LAS/LAZ files to the Jetson.',
            ),
        ],
    ),
    (
        '19. Mission Control User Experience',
        [
            (
                'Start Button',
                'The Start action should build or verify the Jetson workspace, launch the field '
                'ROS stack, wait for readiness, then request the initial reference scan. Mission '
                'status should clearly show whether the system is starting, scanning, navigating, '
                'uploading, processing, or stopped.',
            ),
            (
                'Camera Preview',
                'The bridge can forward a low-rate camera preview so the operator can see current '
                'AI detections and robot context. The preview should be rate-limited to avoid '
                'competing with scan transfers over Wi-Fi.',
            ),
            (
                'Stop and Download',
                'Stop should end the Jetson ROS launch, stop the scan loop, and download compact '
                'digital-twin artifacts. Raw scans should remain on the Perspective host unless '
                'explicitly configured otherwise.',
            ),
        ],
    ),
    (
        '20. Configuration Files',
        [
            (
                'field.env',
                'config/field.env is the operator-specific configuration file. It contains topic '
                'names, paths, robot backend selection, Spot credentials, OAK settings, Trimble '
                'URLs, and digital-twin paths. It is intentionally ignored by git because it can '
                'contain credentials.',
            ),
            (
                'site_calibration.yaml',
                'The site calibration file records calibration values for point-cloud fusion and '
                'site-specific setup. It should be checked carefully whenever the camera, mount, '
                'or payload geometry changes.',
            ),
            (
                'navigation_priorities.yaml',
                'The navigation priority configuration gives the planner a compact way to rank '
                'different defect classes or movement targets. This allows mission behavior to '
                'change without rewriting planner code.',
            ),
        ],
    ),
    (
        '21. ROS Topics and Interfaces',
        [
            (
                'Core Detection Topics',
                'The core AI topics are /oak/rgb/image_raw, /detections_2d, /oak/rgb/depth, '
                '/oak/rgb/camera_info, and /detections_3d. These connect camera input, YOLO '
                'output, depth projection, and marker generation.',
            ),
            (
                'Digital Twin Topics',
                'The digital twin topics include /trimble/x7/scan_points, /digital_twin/map, '
                '/digital_twin/defect_markers, /digital_twin/rescan_goals, /digital_twin/anchor_status, '
                'and /digital_twin/waypoint_arrived.',
            ),
            (
                'Planner and Robot Topics',
                'The planner publishes /infrastructure/inspection_goal and /infrastructure/planner_status. '
                'The robot bridge publishes /infrastructure/navigation_status and waypoint arrival. '
                'These topics separate planning from actual robot command execution.',
            ),
        ],
    ),
    (
        '22. Testing and Simulation',
        [
            (
                'Unit Tests',
                'The repository contains tests for pointcloud normalization, OAK depth fusion, '
                'robot goal bridge arrival behavior, and fusion pipeline utilities. These tests '
                'should be run after edits to launch configuration, message conversion, or robot '
                'goal handling.',
            ),
            (
                'Dry Run Mission Tests',
                'Before field deployment, run the mission loop with ROBOT_GOAL_BACKEND=dry_run. '
                'This verifies scan requests, status updates, digital-twin writes, and UI behavior '
                'without moving Spot.',
            ),
            (
                'Hardware Bringup Tests',
                'Hardware tests should confirm OAK topics, camera preview, OAK odometry, Spot '
                'network reachability, Trimble export folder behavior, scan transfer size, and '
                'arrival-check stability.',
            ),
        ],
    ),
    (
        '23. Deployment Procedure',
        [
            (
                'Preflight',
                'Preflight includes checking power regulation, cable strain relief, OAK USB '
                'connection, router connectivity, Jetson storage, Trimble battery, Perspective '
                'export path, Spot network reachability, and field.env values.',
            ),
            (
                'Startup',
                'Place Spot in the correct starting location, power the payload stack, connect '
                'the control host, start the mission app, and press Start. The system should '
                'launch ROS, verify readiness, and request the initial reference scan.',
            ),
            (
                'Mission Loop',
                'During the mission, the robot navigates to selected inspection goals, confirms '
                'arrival, requests Trimble scans when useful, processes scan files, updates the '
                'digital twin, and repeats until Stop is pressed.',
            ),
            (
                'Shutdown',
                'Stop should terminate ROS cleanly, return or release robot leases when appropriate, '
                'download compact digital-twin outputs, and leave raw scan files on the Perspective '
                'host for later archival.',
            ),
        ],
    ),
    (
        '24. Risks and Mitigations',
        [
            (
                'Localization Drift',
                'OAK VSLAM can drift over time. Use short station-to-station movements, good '
                'lighting, textured views, arrival tolerance checks, and periodic Trimble reference '
                'scans to reduce impact.',
            ),
            (
                'Wireless Transfer Limits',
                'Trimble scans can be large. Keep raw scans on the Perspective host, send reduced '
                'LAS/LAZ copies to the Jetson, and tune max point count before field missions.',
            ),
            (
                'Mechanical Vibration',
                'Loose mounts or cable tugging can degrade camera calibration and depth output. '
                'Use rear M4 mounting, strain relief, vibration-resistant fasteners, and a generous '
                'front optical opening.',
            ),
            (
                'Automation Assumptions',
                'If Perspective does not expose an API or predictable export workflow, keep the '
                'operator in the loop for scan confirmation. The software should still track state '
                'and transfer files after export.',
            ),
        ],
    ),
    (
        '25. Maintenance and Future Work',
        [
            (
                'Immediate Maintenance',
                'Keep field.env updated with actual topic names and IPs, verify model paths, '
                'archive raw Trimble scans, clean OAK lenses and IR windows, and inspect mounts '
                'after transport.',
            ),
            (
                'Recommended Enhancements',
                'Future work includes a native tablet web bridge, better Perspective export '
                'automation, richer mission dashboards, robust scan compression, more formal '
                'digital-twin packaging, and deeper integration with Spot safety workflows.',
            ),
            (
                'Acceptance Criteria',
                'The system is ready for staged field testing when it can start from one button, '
                'confirm OAK topics, create an initial reference scan, publish AI markers, navigate '
                'or dry-run through goals, request gated rescans, and export a compact twin package.',
            ),
        ],
    ),
]


def add_chapter(document, title, sections):
    new_chapter(document, title)
    for subtitle, body in sections:
        document.add_heading(subtitle, level=2)
        para(document, body)
        para(document, section_detail(title, subtitle))


def section_detail(chapter_title, subtitle):
    context = {
        'Inspection Mission Concept': (
            'In deployment, the most important practical behavior is sequencing: the robot '
            'should move only when perception and localization are healthy, and the X7 should '
            'scan only when the robot is stopped and the scan request is intentional. This '
            'separation keeps the fast autonomy loop and the high-resolution survey loop from '
            'fighting each other.'
        ),
        'Primary Data Flow': (
            'The data flow is intentionally staged rather than monolithic. Camera detections '
            'can run continuously, the scan decision node can choose whether a stationary scan '
            'is worth the time, and the digital-twin layer can update after each completed '
            'scan without requiring the entire mission to restart.'
        ),
        'Operational Outcome': (
            'The final field package is meant to be reviewable by a human inspector: compact '
            'YAML marker files preserve AI observations, the anchor file explains the coordinate '
            'relationship, and the full raw scan files remain available on the Perspective host '
            'for post-processing or formal archival.'
        ),
        'Boston Dynamics Spot': (
            'Spot should be treated as a capable but safety-critical mobility platform. The '
            'application can request goals, but payload power, robot lease behavior, operator '
            'authority, and emergency stop practices remain part of the deployment procedure.'
        ),
        'Spot GXP Payload Interface': (
            'The GXP is the clean payload integration point because it avoids improvised wiring '
            'around the robot body. Any use of its regulated power outputs should include current '
            'budgeting for Jetson startup, router draw, OAK USB load, lighting accessories, and '
            'environmental derating.'
        ),
        'Jetson Orin Nano': (
            'The Jetson should be configured as the reliable field computer: fixed hostname or '
            'static DHCP, automatic time configuration where possible, sufficient swap/storage, '
            'and a tested ROS environment. It is also the right place to keep robot credentials '
            'and field-specific configuration outside source control.'
        ),
        'Luxonis OAK-D-W / OAK-D Pro W Class Camera': (
            'Mechanically, the OAK should be bolted from the back and given a large front optical '
            'opening. Electrically, it needs a stable USB connection. Software bringup should '
            'confirm RGB, depth, camera_info, and odometry/VSLAM topics before the mission loop '
            'is trusted.'
        ),
        'Trimble X7': (
            'The X7 is the accuracy anchor for the system, not the continuous navigation sensor. '
            'It should be used at deliberate station stops, where the robot is stable, the scan '
            'export is complete, and the file watcher can safely ingest the finished file.'
        ),
        'Control Host': (
            'The control host is deliberately outside the ROS graph. That makes it easier to run '
            'on a Windows tablet or laptop while the Jetson remains the ROS machine. The common '
            'contract is HTTP endpoints plus scan-file transfer, not ROS installed on Windows.'
        ),
        'Network and Power Hardware': (
            'Network and power failures are among the easiest ways to make a good autonomy stack '
            'look unreliable. The mission LAN should be tested under payload power, with the robot '
            'walking, camera streaming, and scan files transferring at the same time.'
        ),
        'OAK Camera Case': (
            'A good case supports the camera without touching or shadowing the optical path. Use '
            'approximately 3 mm walls, a stronger rear mounting region, small exterior chamfers, '
            'and a larger chamfer around the front opening so the wide stereo cameras do not see '
            'a tunnel.'
        ),
        'Ventilation': (
            'The preferred layout is top exhaust and side intake. Slots should be small enough to '
            'retain strength, placed away from the lenses, and shielded when the system may see '
            'dust, dripping water, or debris.'
        ),
        'Trimble X7 Mounting': (
            'The scanner mount should prioritize repeatability and stillness over compactness. '
            'Any bracket must avoid flex during Spot motion and should be checked after transport '
            'because even small mechanical shifts can affect scan registration assumptions.'
        ),
        'Lighting': (
            'Lighting should improve AI detection without washing out the camera or creating '
            'unwanted reflections. In dark infrastructure spaces, test both RGB detection and '
            'stereo depth with the actual light angle before finalizing the mount.'
        ),
        'Jetson Power': (
            'A barrel-jack supply should be regulated, fused, and physically retained. If the '
            'Jetson browns out, symptoms may appear as camera disconnects, ROS node crashes, or '
            'corrupt scan transfers rather than an obvious power warning.'
        ),
        'Router Power': (
            'The router is small but mission-critical. USB-C power should be secure, and the '
            'router should be mounted where the antenna pattern is not blocked by metal payload '
            'plates or the robot body.'
        ),
        'Camera Power': (
            'The OAK should be fed from a regulated USB-C supply path, similar to the router. '
            'The Jetson connection should be treated as USB data unless the specific carrier '
            'board and cable are intentionally sized for camera power.'
        ),
        'Trimble Power': (
            'Keeping the X7 on its own power domain reduces risk to the robot payload bus. Battery '
            'state should be part of preflight because a scan failure mid-mission can leave the '
            'digital twin incomplete even when the robot stack is healthy.'
        ),
        'Mission LAN': (
            'A private router-based LAN gives predictable addressing and avoids dependency on site '
            'internet. The field app should use IPs or hostnames that are tested before the robot '
            'enters the inspection area.'
        ),
        'ROS and HTTP Boundaries': (
            'The boundary is simple by design: ROS topics stay on the Jetson, while HTTP and file '
            'transfer connect to the Windows control host. This keeps Perspective automation and '
            'Windows file handling out of the robot-side ROS nodes.'
        ),
        'Suggested Addressing': (
            'Document the Jetson, control host, router, and Spot addresses in the field notebook '
            'and in config/field.env. Avoid ad hoc IP changes during a mission because scan requests '
            'and downloads depend on stable routing.'
        ),
    }
    if subtitle in context:
        return context[subtitle]
    if 'Detection' in chapter_title or 'AI' in chapter_title:
        return (
            'For this part of the application, field validation should include a known visual '
            'target, at least one negative scene, and one low-light scene. The goal is to verify '
            'both the model output and the downstream behavior triggered by that output.'
        )
    if 'Trimble' in chapter_title or 'Scan' in chapter_title:
        return (
            'The important deployment check is file completeness. The software should ingest only '
            'stable exported files, not partially written scan data, and the operator should know '
            'where raw and reduced scans are stored after each station.'
        )
    if 'Robot' in chapter_title or 'Motion' in chapter_title or 'Arrival' in chapter_title:
        return (
            'Before enabling physical motion, run this section with the dry-run backend and inspect '
            'the status topics. Live movement should only be enabled after localization, transforms, '
            'leases, and emergency stop procedures are verified.'
        )
    if 'Configuration' in chapter_title or 'ROS Topics' in chapter_title:
        return (
            'Configuration should be changed in one place, recorded before field use, and checked '
            'with ros2 topic list, ros2 topic info, and a short dry-run mission. Topic mismatches '
            'are one of the most common bringup problems.'
        )
    if 'Testing' in chapter_title or 'Deployment' in chapter_title:
        return (
            'This section should be treated as an operational checklist rather than background '
            'reading. Each item maps to a failure mode that is much cheaper to catch in staging '
            'than during a live inspection.'
        )
    return (
        'This component should be reviewed in the context of the full mission loop: perception, '
        'localization, scan request, file transfer, digital-twin update, goal selection, and safe '
        'shutdown. Its value is highest when its interfaces are tested with the adjacent components.'
    )


def add_appendices(document):
    new_chapter(document, 'Appendix A. Configuration Snapshot')
    code_block(
        document,
        '\n'.join(
            [
                'IMAGE_TOPIC=/oak/rgb/image_raw',
                'OAK_DEPTH_TOPIC=/oak/rgb/depth',
                'OAK_CAMERA_INFO_TOPIC=/oak/rgb/camera_info',
                'OAK_ODOM_TOPIC=/oak/odom',
                'ROBOT_WORLD_FRAME=oak_odom',
                'NAVIGATION_BASE_FRAME=body',
                'TRIMBLE_SCAN_TOPIC=/trimble/x7/scan_points',
                'TRIMBLE_WINDOWS_URL=http://CONTROL_HOST_IP:8765',
                'ROBOT_GOAL_BACKEND=dry_run|nav2|http|spot_sdk',
            ]
        ),
    )
    para(
        document,
        'This snapshot is representative. The actual field deployment should use '
        'config/field.env so credentials, IP addresses, and site-specific paths are not committed.',
    )

    new_chapter(document, 'Appendix B. Hardware Checklist')
    bullets(
        document,
        [
            'Spot robot with payload mounting rails and configured operator safety process.',
            'Spot GXP payload breakout for regulated power and communications.',
            'NVIDIA Jetson Orin Nano with barrel-jack power supply path.',
            'Luxonis OAK-D-W / OAK-D Pro W class camera with USB cable and rear M4 mounting.',
            'Trimble X7 scanner with battery and Perspective control host.',
            'Windows tablet or Windows laptop for Perspective control and mission UI.',
            'Mission router or private LAN equipment, preferably powered from a stable regulated supply.',
            'USB-C router power cable, USB-C OAK power cable, Jetson barrel cable, OAK USB data cable, cable glands, strain relief, and fasteners.',
            'Ulanzi 40 W Pro LED or equivalent lighting hardware if low-light inspections are expected.',
            '3D-printed or machined OAK case with 3 mm walls, 5 mm mount region, vents, and a large chamfered optical opening.',
        ],
    )

    new_chapter(document, 'Appendix C. Field Acceptance Checklist')
    numbered(
        document,
        [
            'Confirm all devices join the mission LAN and have known IP addresses.',
            'Confirm the Jetson can see OAK RGB, depth, camera_info, and odometry topics.',
            'Confirm YOLO publishes /detections_2d on real camera input.',
            'Confirm OAK depth fusion publishes /detections_3d for known targets.',
            'Confirm the Perspective host receives /scan_request and /waypoint_arrived events.',
            'Confirm a completed LAS or LAZ export is transferred to the Jetson scan folder.',
            'Confirm /trimble/x7/scan_points and /digital_twin/map are published after scan ingest.',
            'Confirm /tmp/digital_twin_defects.yaml and /tmp/digital_twin_anchor.yaml are created.',
            'Confirm robot_goal_bridge is still dry_run before enabling physical motion.',
            'Confirm Spot command backend, credentials, lease behavior, and emergency stop procedure before live movement.',
        ],
    )


def add_implementation_deep_dive(document):
    new_chapter(document, '26. Python Package and Entry Point Implementation')
    document.add_heading('ROS Package Boundaries', level=2)
    para(
        document,
        'The Python implementation is split into two ROS packages. The defect_detection '
        'package contains the application logic: camera publishing, YOLO inference, OAK '
        'depth fusion, scan gating, digital-twin state, planning, visualization, and robot '
        'goal command dispatch. The pointcloud_bridge package is deliberately smaller and '
        'only normalizes incoming PointCloud2 streams. Keeping pointcloud_bridge separate '
        'prevents the generic cloud transport code from depending on the larger inspection '
        'application.'
    )
    add_table(
        document,
        ['Package', 'Primary Python Modules', 'Role'],
        [
            ['defect_detection', 'defect_detection.defect_detection.*', 'YOLO detection, point-cloud fusion, and visualization nodes.'],
            ['defect_detection', 'defect_detection.digital_twin.*', 'Trimble scan ingest, map building, defect persistence, planning, and robot bridge nodes.'],
            ['defect_detection', 'defect_detection.autonomous_navigation.navigator', 'Older direct defect navigation logic using 3D detections and Nav2.'],
            ['defect_detection', 'defect_detection.spot_cam_loading.*', 'OpenCV camera publisher and simple subscriber utilities.'],
            ['pointcloud_bridge', 'pointcloud_bridge.pointcloud_bridge', 'Timestamp and field validation for PointCloud2 transport.'],
            ['tools', 'tools/trimble_perspective_bridge/windows_app.py', 'Control host UI, HTTP server, scan export watcher, SSH/SCP coordination.'],
        ],
    )
    document.add_heading('Console Script Mapping', level=2)
    para(
        document,
        'setup.py exposes each ROS node as a console script. Launch files call these script '
        'names rather than importing Python modules directly. This is important because the '
        'same node can be started from ros2 launch, manually with ros2 run, or during tests '
        'after the package is installed with colcon.'
    )
    add_table(
        document,
        ['Console Script', 'Python Class or main()', 'Implementation Purpose'],
        [
            ['yolo_detector', 'YoloNode', 'Loads the trained model and publishes Detection2DArray messages.'],
            ['oak_depth_fusion_node', 'OakDepthFusionNode', 'Combines YOLO 2D detections with aligned OAK depth into Detection3DArray.'],
            ['oak_localization_bridge', 'OakLocalizationBridge', 'Republishes OAK odometry as TF for robot-localized planning.'],
            ['scan_decision_node', 'ScanDecisionNode', 'Decides when detections justify an X7 scan request.'],
            ['trimble_windows_bridge', 'TrimbleWindowsBridge', 'Posts scan and status events from ROS to the Perspective control host.'],
            ['trimble_scan_watcher', 'TrimbleScanWatcher', 'Converts stable LAS/LAZ exports into /trimble/x7/scan_points.'],
            ['pointcloud_to_occupancy', 'PointCloudToOccupancy', 'Converts Trimble point clouds into a planning occupancy grid.'],
            ['defect_map_node', 'DefectMapNode', 'Persists AI markers and publishes rescan goals.'],
            ['infrastructure_planner', 'InfrastructurePlanner', 'Chooses defect rescan or frontier inspection goals.'],
            ['robot_goal_bridge', 'RobotGoalBridge', 'Dispatches inspection goals to dry_run, Nav2, HTTP, or Spot SDK backends.'],
        ],
    )
    code_block(
        document,
        "entry_points={\n"
        "    'console_scripts': [\n"
        "        'oak_depth_fusion_node = defect_detection.digital_twin.oak_depth_fusion_node:main',\n"
        "        'robot_goal_bridge = defect_detection.digital_twin.robot_goal_bridge:main',\n"
        "        'trimble_scan_watcher = defect_detection.digital_twin.trimble_scan_watcher:main',\n"
        "    ]\n"
        "}"
    )

    new_chapter(document, '27. Launch and Runtime Configuration Flow')
    document.add_heading('Field Script Startup Path', level=2)
    para(
        document,
        'scripts/run_field.sh is the main Jetson startup wrapper. It loads config/field.env, '
        'sources ROS 2 Jazzy and the workspace install setup, runs the field preflight, then '
        'calls ros2 launch pointcloud_bridge full_pipeline.launch.xml. This lets the control '
        'host start the whole ROS graph with one command while keeping site-specific values in '
        'a single environment file.'
    )
    numbered(
        document,
        [
            'Resolve WORKSPACE_ROOT from the script location.',
            'Load FIELD_CONFIG, normally config/field.env.',
            'Source /opt/ros/jazzy/setup.bash and install/setup.bash.',
            'Create runtime log and YOLO config directories.',
            'Run scripts/field_preflight.sh.',
            'Compute detector/fusion flags based on transport or full mode.',
            'Launch pointcloud_bridge full_pipeline.launch.xml with environment-backed arguments.',
        ],
    )
    document.add_heading('Launch File Layering', level=2)
    para(
        document,
        'The top-level launch file lives in the pointcloud_bridge package because it includes '
        'both the generic point cloud bridge and the larger defect_detection launch file. The '
        'defect_detection launch file then instantiates individual ROS nodes based on Boolean '
        'launch arguments. This keeps the full pipeline composable: transport-only runs can '
        'skip detector and planner nodes, while full field runs can enable the entire stack.'
    )
    add_table(
        document,
        ['Launch Argument', 'Default/Example', 'Effect'],
        [
            ['detector', 'true in full mode', 'Starts yolo_detector and publishes /detections_2d.'],
            ['oak_depth_navigation', 'OAK_DEPTH_NAVIGATION=true', 'Starts oak_depth_fusion_node for 2D-to-3D conversion.'],
            ['trimble_scan_watcher', 'true', 'Starts the LAS/LAZ folder watcher.'],
            ['scan_decision', 'true', 'Starts high-confidence scan request gating.'],
            ['digital_twin_map', 'true', 'Starts pointcloud_to_occupancy.'],
            ['infrastructure_planner', 'true', 'Starts the planner that publishes /infrastructure/inspection_goal.'],
            ['robot_goal_bridge', 'false by default', 'Starts the motion backend bridge only when explicitly enabled.'],
        ],
    )
    document.add_heading('Configuration Ownership', level=2)
    para(
        document,
        'The field.env file owns deploy-time values: IP addresses, topic names, model paths, '
        'Spot credentials, scan folders, and backend choice. Python nodes declare safe defaults, '
        'but the launch layer overrides them from field.env so the code does not need to be edited '
        'for each site.'
    )
    code_block(
        document,
        "OAK_DEPTH_NAVIGATION=true\n"
        "OAK_DEPTH_TOPIC=/oak/rgb/depth\n"
        "OAK_CAMERA_INFO_TOPIC=/oak/rgb/camera_info\n"
        "TRIMBLE_WINDOWS_URL=http://CONTROL_HOST_IP:8765\n"
        "ROBOT_GOAL_BACKEND=dry_run\n"
        "ROBOT_WORLD_FRAME=oak_odom"
    )

    new_chapter(document, '28. OAK Depth Fusion Python Internals')
    document.add_heading('Class Responsibility', level=2)
    para(
        document,
        'OakDepthFusionNode is a small synchronous fusion node. It does not run neural inference '
        'itself. Instead, it consumes Detection2DArray messages from the YOLO node, consumes the '
        'aligned OAK depth image, waits for camera_info intrinsics, and publishes Detection3DArray. '
        'This separation makes it possible to test depth projection logic independently from model '
        'loading and GPU inference.'
    )
    document.add_heading('Important Functions', level=2)
    add_table(
        document,
        ['Function', 'Input', 'Output/Effect'],
        [
            ['detection_confidence()', 'vision_msgs Detection2D', 'Extracts first hypothesis score or returns 0.0.'],
            ['detection_class_id()', 'vision_msgs Detection2D', 'Extracts first class id or returns unknown.'],
            ['bbox_bounds()', 'Detection bbox and image size', 'Computes clipped integer ROI bounds with optional padding.'],
            ['depth_image_to_meters()', 'CV depth array and encoding', 'Normalizes 16UC1 millimeters or 32FC1 meters into float meters.'],
            ['robust_depth()', 'Depth ROI', 'Filters invalid values and returns a percentile-clipped median depth.'],
            ['project_pixel_to_camera()', 'u, v, depth, K matrix', 'Projects detection center into camera x, y, z coordinates.'],
            ['convert_detection()', 'Detection2D and depth image', 'Builds one Detection3D with pose and bbox.'],
        ],
    )
    document.add_heading('Message Synchronization', level=2)
    para(
        document,
        'message_filters.ApproximateTimeSynchronizer is used because RGB detections and depth '
        'images rarely have exactly identical timestamps. The node accepts messages within the '
        'configured slop window, defaulting to 0.12 seconds. This is a practical compromise: too '
        'tight a slop loses valid pairs; too loose a slop can associate detections with stale depth.'
    )
    code_block(
        document,
        "self.synchronizer = message_filters.ApproximateTimeSynchronizer(\n"
        "    [self.detections_sub, self.depth_sub],\n"
        "    queue_size=sync_queue_size,\n"
        "    slop=sync_slop_sec,\n"
        ")\n"
        "self.synchronizer.registerCallback(self.synchronized_callback)"
    )
    document.add_heading('Depth Robustness', level=2)
    para(
        document,
        'The robust_depth function intentionally avoids taking a single depth sample at the box '
        'center. Stereo depth can have holes or edge noise, especially around defects, cracks, or '
        'reflective surfaces. The implementation samples the padded detection region, removes '
        'invalid and very small values, clips to the 15th through 85th percentile, and returns the '
        'median. This produces a stable representative depth while rejecting outliers.'
    )

    new_chapter(document, '29. OAK Localization Bridge Internals')
    document.add_heading('Odometry to TF Conversion', level=2)
    para(
        document,
        'OakLocalizationBridge subscribes to nav_msgs/Odometry on /oak/odom and broadcasts a TF '
        'transform. This gives the rest of the system a standard transform tree, even if the OAK '
        'odometry source is produced by DepthAI ROS, RTAB-Map, or another VSLAM node.'
    )
    add_table(
        document,
        ['Parameter', 'Default', 'Implementation Meaning'],
        [
            ['odom_topic', '/oak/odom', 'Input odometry topic.'],
            ['odom_frame', 'oak_odom', 'Default parent frame when message frame ids are ignored.'],
            ['base_frame', 'base_link', 'Default child frame. Field config sets this to body.'],
            ['use_message_frame_ids', 'true', 'When true, trust header.frame_id and child_frame_id from odometry.'],
            ['zero_z', 'true', 'Flatten z translation to 0 for 2D planning if desired.'],
        ],
    )
    document.add_heading('Why TF Is Central', level=2)
    para(
        document,
        'The planner, frame anchor, and robot goal bridge all depend on transforms. The OAK odometry '
        'bridge is what lets a PoseStamped goal in map or oak_odom frame be compared against the '
        'current robot body pose. If this transform is missing, arrival checking and coordinate '
        'conversion will fail even if camera images and detections are publishing normally.'
    )

    new_chapter(document, '30. Scan Decision Node Internals')
    document.add_heading('State Machine', level=2)
    para(
        document,
        'ScanDecisionNode is a timer-driven gate. It stores the most recent Detection2DArray, tracks '
        'when it was received, and periodically decides whether the latest detections are fresh, '
        'numerous enough, confident enough, and outside the cooldown window. It publishes both a '
        'Bool scan request and a String reason so the control host can display why a scan was requested.'
    )
    add_table(
        document,
        ['Rule', 'Default', 'Purpose'],
        [
            ['confidence_threshold', '0.65', 'Ignore uncertain detections.'],
            ['min_detections', '1', 'Require at least one qualifying detection.'],
            ['detection_timeout_sec', '3.0', 'Ignore stale detections.'],
            ['scan_cooldown_sec', '60.0', 'Avoid repeated scans from the same short-lived observation.'],
            ['decision_period_sec', '2.0', 'Periodic scan gate evaluation rate.'],
        ],
    )
    document.add_heading('Why This Exists', level=2)
    para(
        document,
        'Trimble scans are valuable but slow compared with camera inference. Without a gate, the '
        'system could request a scan for every frame or every low-confidence false positive. The '
        'scan decision node makes scan requests sparse, explainable, and tunable from launch parameters.'
    )

    new_chapter(document, '31. Trimble Bridge Python Internals')
    document.add_heading('ROS-to-HTTP Adapter', level=2)
    para(
        document,
        'TrimbleWindowsBridge is a ROS node that translates ROS events into HTTP calls for the '
        'Windows Perspective control host. The node '
        'posts readiness, mission status, scan requests, waypoint arrival events, inspection goals, '
        'navigation status, detections, and optional low-rate camera preview frames.'
    )
    add_table(
        document,
        ['ROS Input', 'Callback', 'HTTP/State Effect'],
        [
            ['/digital_twin/scan_required', 'scan_required_callback', 'POST /scan_request when true and not suppressed.'],
            ['/digital_twin/scan_reason', 'scan_reason_callback', 'Stores human-readable scan reason.'],
            ['/digital_twin/waypoint_arrived', 'waypoint_arrived_callback', 'POST /waypoint_arrived and trigger scan status.'],
            ['/digital_twin/frontier_goal', 'frontier_goal_callback', 'Stores latest frontier goal for context.'],
            ['/infrastructure/inspection_goal', 'inspection_goal_callback', 'POST planner goal metadata to control host.'],
            ['/infrastructure/navigation_status', 'navigation_status_callback', 'POST mission status updates.'],
            ['/detections_2d', 'detections_callback', 'Stores compact detection summary for UI.'],
            ['image_topic', 'image_callback', 'Encodes low-rate JPEG preview and posts it.'],
        ],
    )
    document.add_heading('HTTP Payload Handling', level=2)
    para(
        document,
        'The post_json helper uses urllib.request rather than a large web client dependency. It '
        'serializes dictionaries to JSON, sets content-type, applies a timeout, and returns HTTP '
        'status and response body. Errors are caught around call sites so the ROS graph can keep '
        'running even if the control host is temporarily unreachable.'
    )
    code_block(
        document,
        "def post_json(url, payload, timeout_sec):\n"
        "    data = json.dumps(payload).encode('utf-8')\n"
        "    request = urlrequest.Request(url, data=data,\n"
        "        headers={'content-type': 'application/json'}, method='POST')\n"
        "    with urlrequest.urlopen(request, timeout=timeout_sec) as response:\n"
        "        return response.status, response.read().decode('utf-8')"
    )
    document.add_heading('Camera Preview Implementation', level=2)
    para(
        document,
        'Preview is optional because image transfer can consume network bandwidth. When enabled, '
        'the node uses cv_bridge and OpenCV to decode the ROS image, resize it to a configured '
        'width, overlay or attach recent detection context, encode it as JPEG, base64 encode the '
        'result, and send it to the control host. The preview is rate-limited by camera_preview_rate_hz.'
    )

    new_chapter(document, '32. Trimble Scan Watcher Internals')
    document.add_heading('File Stability Logic', level=2)
    para(
        document,
        'TrimbleScanWatcher does not publish a file as soon as it appears. It polls the scan folder, '
        'checks file extension, checks modification age, and only ingests files that appear stable. '
        'This prevents partially written LAS/LAZ exports from entering the ROS graph.'
    )
    add_table(
        document,
        ['Parameter', 'Default', 'Effect'],
        [
            ['scan_directory', '/tmp/trimble_scans', 'Folder watched for completed scan exports.'],
            ['output_topic', '/trimble/x7/scan_points', 'PointCloud2 topic published after ingest.'],
            ['frame_id', 'map', 'Frame assigned to scan points.'],
            ['poll_period_sec', '2.0', 'Folder polling interval.'],
            ['stable_age_sec', '5.0', 'Minimum age before file is considered complete.'],
            ['max_points', '500000', 'Point cap for Jetson-friendly scan publication.'],
            ['require_scan_request', 'false', 'When true, only ingest after gate opens.'],
        ],
    )
    document.add_heading('PointCloud2 Construction', level=2)
    para(
        document,
        'After LAS/LAZ points are loaded, the watcher builds a ROS PointCloud2 message with x, '
        'y, and z fields. The max_points parameter is a practical field setting: it lets the '
        'Jetson operate on a manageable cloud while raw full-resolution scan files remain on the '
        'Perspective host for later survey processing.'
    )

    new_chapter(document, '33. Occupancy Map Implementation')
    document.add_heading('Cloud Filtering', level=2)
    para(
        document,
        'PointCloudToOccupancy subscribes to /trimble/x7/scan_points, reads xyz points, filters '
        'them by z range and max range, then projects the remaining points into a 2D occupancy '
        'grid. This is intentionally a planning abstraction rather than a full 3D reconstruction.'
    )
    add_table(
        document,
        ['Parameter', 'Default', 'Reason'],
        [
            ['resolution', '0.10', 'Grid cell size in meters.'],
            ['padding_m', '2.0', 'Extra map border around observed points.'],
            ['min_z', '-0.25', 'Reject floor/below-origin noise.'],
            ['max_z', '1.20', 'Reject high points for 2D navigation planning.'],
            ['max_range_m', '80.0', 'Limit distant points that may not matter to station planning.'],
            ['scan_origin_x/y', '0.0', 'Reference origin used during grid construction.'],
        ],
    )
    document.add_heading('Map Publication', level=2)
    para(
        document,
        'The output OccupancyGrid is published on /digital_twin/map. Other nodes treat this as '
        'the planning representation: frontier_planner can search its boundaries, infrastructure_planner '
        'can use it when no defect rescan goals exist, and RViz can visualize it as a compact map.'
    )

    new_chapter(document, '34. Defect Map Node Implementation')
    document.add_heading('Merging and Persistence', level=2)
    para(
        document,
        'DefectMapNode receives Detection3DArray messages and maintains a dictionary of observed '
        'defects. New detections are either merged into an existing defect within merge_radius_m '
        'or added as a new defect. The node writes YAML state so detections survive beyond the '
        'lifetime of a single ROS message.'
    )
    document.add_heading('Marker and Rescan Outputs', level=2)
    add_table(
        document,
        ['Output', 'Message Type', 'Use'],
        [
            ['/digital_twin/defect_markers', 'visualization_msgs/MarkerArray', 'RViz and mission-control visualization.'],
            ['/digital_twin/rescan_goals', 'geometry_msgs/PoseArray or similar goal representation', 'Planner input for revisiting defects.'],
            ['/tmp/digital_twin_defects.yaml', 'YAML file', 'Compact final digital-twin artifact.'],
        ],
    )
    para(
        document,
        'The implementation favors simple persistence over a database because the field artifact '
        'needs to be easy to copy, inspect, and archive. YAML is sufficient for defect class, pose, '
        'confidence, observation count, and timestamp-style metadata.'
    )

    new_chapter(document, '35. Infrastructure Planner Implementation')
    document.add_heading('Goal Selection Order', level=2)
    para(
        document,
        'InfrastructurePlanner is the high-level station selector. It periodically checks whether '
        'planning is enabled and whether it is outside the cooldown window. It prefers defect rescan '
        'goals because those represent known AI findings. If no defect goals are pending, it falls '
        'back to map frontier or general exploration behavior.'
    )
    numbered(
        document,
        [
            'Read latest map and latest rescan goal state.',
            'Skip planning when disabled or within cooldown.',
            'Choose a defect rescan station when one is available and preferred.',
            'Otherwise compute a frontier or map-edge station.',
            'Publish a PoseStamped to /infrastructure/inspection_goal.',
            'Publish planner status text for the control host and logs.',
        ],
    )
    document.add_heading('Coordinate Frames', level=2)
    para(
        document,
        'Planner output is expressed in a target frame, normally map. The robot_goal_bridge later '
        'transforms this goal into the command frame expected by the selected backend. This split '
        'keeps station selection independent from whether movement is dry-run, Nav2, HTTP, or '
        'Spot SDK controlled.'
    )

    new_chapter(document, '36. Robot Goal Bridge Implementation')
    document.add_heading('Backend Dispatch', level=2)
    para(
        document,
        'RobotGoalBridge subscribes to /infrastructure/inspection_goal and stores one active goal '
        'at a time. If a new goal arrives while another is active, it publishes a busy status and '
        'ignores the new goal. This prevents overlapping robot commands.'
    )
    add_table(
        document,
        ['Backend', 'Method', 'Implementation Behavior'],
        [
            ['dry_run', 'create_timer()', 'No robot command; publishes arrival after dry_run_arrival_delay_sec.'],
            ['nav2', 'ActionClient NavigateToPose', 'Sends goal to /navigate_to_pose and waits for action result.'],
            ['http', 'post_json()', 'POSTs pose payload to an external command service.'],
            ['spot_sdk', 'bosdyn.client RobotCommandClient', 'Connects to Spot, leases, stands, sends SE2 trajectory command.'],
        ],
    )
    document.add_heading('Arrival Verification State', level=2)
    para(
        document,
        'The bridge stores active_goal, active_goal_started_sec, arrival_stable_since, and an '
        'arrival_timer. When TF arrival checking is enabled, it repeatedly transforms the current '
        'robot body pose into the goal frame and compares distance and yaw against tolerance. '
        'The robot must remain inside tolerance for arrival_stable_sec before waypoint arrival '
        'is published.'
    )
    code_block(
        document,
        "self.active_goal = goal\n"
        "self.active_goal_started_sec = time.monotonic()\n"
        "self.arrival_stable_since = None\n"
        "...\n"
        "if self.backend == 'nav2':\n"
        "    self.send_nav2_goal(goal)\n"
        "elif self.backend == 'http':\n"
        "    self.send_http_goal(goal)\n"
        "elif self.backend == 'spot_sdk':\n"
        "    self.send_spot_sdk_goal(goal)\n"
        "else:\n"
        "    self.dry_run_timer = self.create_timer(...)"
    )
    document.add_heading('Spot SDK Connection', level=2)
    para(
        document,
        'The Spot SDK path lazily connects only when needed. ensure_spot_connected imports the '
        'Bosdyn SDK, reads credentials from parameters or environment variables, authenticates, '
        'waits for time sync, acquires a lease keepalive, creates RobotCommandClient, optionally '
        'powers on motors, and optionally commands stand. This keeps startup safe when the backend '
        'is not set to spot_sdk.'
    )

    new_chapter(document, '37. Control Host Application Implementation')
    document.add_heading('Windows App Responsibilities', level=2)
    para(
        document,
        'windows_app.py is both a GUI and a small mission server. It maintains operator-editable '
        'configuration fields, launches remote Jetson commands over SSH, listens for HTTP callbacks '
        'from ROS, watches the Perspective export folder, prepares reduced scan files, transfers '
        'them to the Jetson, and downloads mission artifacts at stop.'
    )
    add_table(
        document,
        ['HTTP Endpoint', 'Called By', 'Purpose'],
        [
            ['POST /jetson_ready', 'TrimbleWindowsBridge', 'Tell UI that ROS stack started.'],
            ['POST /scan_request', 'TrimbleWindowsBridge', 'Ask host/Perspective workflow for a scan.'],
            ['POST /waypoint_arrived', 'RobotGoalBridge through bridge node', 'Notify host that robot is at station.'],
            ['GET /health', 'ROS or operator tooling', 'Verify bridge server is reachable.'],
            ['GET /status', 'Browser UI or app internals', 'Return current mission state.'],
            ['GET /camera_frame', 'Browser UI', 'Return latest preview image.'],
        ],
    )
    document.add_heading('Windows Control Host Deployment', level=2)
    para(
        document,
        'The field control host is the Windows tablet or laptop running Trimble Perspective and '
        'tools/trimble_perspective_bridge/windows_app.py. The Windows installer batch installs '
        'Paramiko, watchdog, laspy, lazrs, numpy, and requests. The launcher batch starts the '
        'Tkinter console, which exposes port 8765 for Jetson status and scan requests.'
    )

    new_chapter(document, '38. Error Handling and Observability')
    document.add_heading('Status Topics', level=2)
    para(
        document,
        'Most nodes publish human-readable String status messages in addition to machine-readable '
        'topics. This is intentional because field debugging often happens from a terminal, tablet, '
        'or RViz session where a concise reason is more useful than a raw exception.'
    )
    add_table(
        document,
        ['Component', 'Status Channel', 'Example Message'],
        [
            ['scan_decision_node', '/digital_twin/scan_reason', 'high-confidence detection'],
            ['frame_anchor_node', '/digital_twin/anchor_status', 'anchored map to oak_odom'],
            ['infrastructure_planner', '/infrastructure/planner_status', 'publishing defect rescan goal'],
            ['robot_goal_bridge', '/infrastructure/navigation_status', 'spot command sent / arrived / failed'],
            ['trimble_windows_bridge', 'HTTP /status updates', 'Scanning, Navigating, Uploading, Jetson Ready'],
        ],
    )
    document.add_heading('Graceful Degradation', level=2)
    para(
        document,
        'The implementation tries to keep the graph alive when noncritical connections fail. For '
        'example, failed HTTP posts to the control host are logged rather than crashing the ROS '
        'node. Missing cv_bridge disables camera preview but does not disable scan requests. This '
        'behavior is important because field networks and optional UI features are less reliable '
        'than local ROS node execution.'
    )

    new_chapter(document, '39. Testing Strategy for Python Nodes')
    document.add_heading('Existing Test Focus', level=2)
    para(
        document,
        'The current tests focus on deterministic logic: point cloud timestamp validation, OAK '
        'depth projection helpers, robot goal bridge arrival behavior, and fusion pipeline pieces. '
        'These are good test targets because they do not require real Spot hardware or a live X7.'
    )
    add_table(
        document,
        ['Test Area', 'Representative File', 'What It Protects'],
        [
            ['Point cloud bridge', 'src/pointcloud_bridge/test/test_pointcloud_bridge.py', 'Timestamp and required xyz field validation.'],
            ['OAK depth fusion', 'src/defect_detection/test/test_oak_depth_fusion_node.py', 'Depth conversion, robust depth, bbox clipping, projection.'],
            ['Robot goal bridge', 'src/defect_detection/test/test_robot_goal_bridge.py', 'Backend dispatch and arrival verification behavior.'],
            ['Fusion pipeline', 'src/defect_detection/test/test_fusion_pipeline.py', 'Detection and point-cloud fusion assumptions.'],
        ],
    )
    document.add_heading('Recommended Additional Tests', level=2)
    numbered(
        document,
        [
            'Mock Trimble control host and assert exact /scan_request payloads.',
            'Feed synthetic Detection2DArray and Image messages through oak_depth_fusion_node.',
            'Simulate stale LAS/LAZ files and verify trimble_scan_watcher ignores them until stable.',
            'Mock TF transforms and verify robot_goal_bridge stable-arrival timing.',
            'Run a dry-run full launch with detector disabled and verify status topics appear.',
            'Add a Windows bridge HTTP contract test so the Jetson-to-control-host payloads remain stable.',
        ],
    )

    new_chapter(document, '40. Implementation Maintenance Notes')
    document.add_heading('Where to Change Behavior', level=2)
    para(
        document,
        'Most mission behavior should be changed through configuration, not code. Topic names, '
        'thresholds, backend selection, scan folders, and IPs belong in config/field.env. Code '
        'changes are appropriate when message contracts, planner algorithms, file parsing, or '
        'robot command semantics need to change.'
    )
    add_table(
        document,
        ['Need', 'Preferred Edit Location', 'Why'],
        [
            ['Change OAK topic names', 'config/field.env', 'Different DepthAI launch files can publish different names.'],
            ['Change scan threshold', 'SCAN_CONFIDENCE_THRESHOLD in field.env', 'Field-tunable without rebuild.'],
            ['Change robot backend', 'ROBOT_GOAL_BACKEND in field.env', 'Switch dry-run/nav2/http/spot safely.'],
            ['Change X7 export folder', 'Control host app config', 'The host owns Perspective filesystem access.'],
            ['Change 2D-to-3D math', 'oak_depth_fusion_node.py', 'Projection algorithm is implementation logic.'],
            ['Change station selection algorithm', 'infrastructure_planner.py', 'Planner owns goal priority and cooldown.'],
            ['Change Spot command semantics', 'robot_goal_bridge.py', 'Backend owns lease and command behavior.'],
        ],
    )
    document.add_heading('Code Review Priorities', level=2)
    para(
        document,
        'Future reviews should prioritize safety and mission correctness: no overlapping active '
        'goals, no scans before arrival, no unbounded file transfers, no direct robot motion when '
        'dry-run is expected, and no credentials committed to source. Performance matters, but '
        'predictable state transitions matter more for field deployment.'
    )


def main():
    document = Document()
    configure_document(document)
    add_title_page(document)
    for title, sections in CHAPTERS:
        add_chapter(document, title, sections)
    add_implementation_deep_dive(document)
    add_appendices(document)

    for section in document.sections:
        add_footer(section)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == '__main__':
    main()
